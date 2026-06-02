import streamlit as st
import pandas as pd
from sqlalchemy import create_engine
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import json
from datetime import datetime

# --- KONEKSI DATABASE ---
@st.cache_resource
def get_engine():
    return create_engine(st.secrets["DB_URL"], pool_size=5, max_overflow=10)

engine = get_engine()
# -------------------------------------------------------------

def format_rupiah(x):
    try: return f"{float(x):,.0f}".replace(',', '.')
    except (ValueError, TypeError): return x

def format_tgl_indo(tgl_str):
    if not tgl_str: return ""
    try:
        tgl_clean = str(tgl_str)[:10]
        dt = datetime.strptime(tgl_clean, "%Y-%m-%d")
        bulan_indo = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni", 
                      "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
        return f"{dt.day} {bulan_indo[dt.month]} {dt.year}"
    except:
        return str(tgl_str)[:10]

def split_kode(teks):
    s = str(teks).strip()
    if " - " in s:
        parts = s.split(" - ", 1)
        return parts[0].strip(), parts[1].strip()
    parts = s.split(" ", 1)
    if len(parts) == 2:
        first_part = parts[0].strip()
        if any(c.isdigit() for c in first_part) or len(first_part) <= 8 or "." in first_part:
            return first_part, parts[1].strip()
    if any(c.isdigit() for c in s) or len(s) <= 8 or "." in s:
        return s, ""
    return "", s

@st.cache_data(ttl=60)
def load_active_rab():
    conn = engine.connect()
    try:
        df_utama = pd.read_sql("SELECT * FROM rab_utama WHERE \"Is_Active\" = 1", conn)
        if not df_utama.empty:
            ids = tuple(df_utama['ID_RAB'].tolist())
            if len(ids) == 1:
                df_detail = pd.read_sql(f"SELECT * FROM rab_detail WHERE \"ID_RAB\" = '{ids[0]}'", conn)
            else:
                df_detail = pd.read_sql(f"SELECT * FROM rab_detail WHERE \"ID_RAB\" IN {ids}", conn)
        else:
            df_detail = pd.DataFrame()
    except Exception as e:
        df_utama, df_detail = pd.DataFrame(), pd.DataFrame()
    conn.close()
    return df_utama, df_detail

# --- FUNGSI AI GEMINI (JSON MODE) ---
def generate_narasi_tor_json(kegiatan, total_anggaran, sasaran, list_belanja, poin_tambahan):
    prompt = f"""
    Anda adalah perencana anggaran ahli di Fakultas Ilmu Budaya Universitas Mulawarman. 
    Tugas Anda adalah menulis komponen isi untuk Term of Reference (TOR) berdasarkan data:
    - Kegiatan: {kegiatan}, Sasaran: {sasaran}, Dana: {total_anggaran}, Item: {list_belanja}.
    - Catatan: {poin_tambahan}
    
    ATURAN PENULISAN:
    1. "dasar_hukum": Kembalikan dalam bentuk ARRAY JSON (List) berisi 4-6 string dasar hukum.
    2. "gambaran_umum": Tuliskan dalam SATU paragraf yang komprehensif. Integrasikan Peraturan Menteri terkait.
    3. "penerima_manfaat", "metode_pelaksanaan", "tahapan_waktu": Masing-masing SATU PARAGRAF detail.
    4. "biaya_diperlukan": Tulis SATU PARAGRAF mendetail bahwa total anggaran adalah Rp {total_anggaran} dari dana BOPTN/PNBP FIB Unmul, tanpa tabel.
    
    Kembalikan output JSON murni (tanpa markdown) dengan kunci: 
    "dasar_hukum", "gambaran_umum", "penerima_manfaat", "metode_pelaksanaan", "tahapan_waktu", "biaya_diperlukan".
    """

    try:
        genai.configure(api_key=st.secrets["GEMINI_API_KEY_NEW"])
        daftar_model = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        
        if not daftar_model: return None
        
        model_terpilih = daftar_model[0]
        for m in daftar_model:
            if 'flash' in m:
                model_terpilih = m
                break
            elif 'pro' in m and 'vision' not in m:
                model_terpilih = m
                
        nama_bersih = model_terpilih.replace("models/", "")
        try:
            model = genai.GenerativeModel(nama_bersih)
            respons = model.generate_content(prompt)
        except:
            model = genai.GenerativeModel(model_terpilih)
            respons = model.generate_content(prompt)
            
        teks_respons = respons.text.replace('```json', '').replace('```', '').strip()
        return json.loads(teks_respons)
            
    except Exception as e:
        st.error(f"❌ Gagal total menghubungi server AI. Detail Error: {e}")
        return None

# --- BUILDER MICROSOFT WORD (.DOCX) ---
def build_docx(meta, narasi, tampilkan_paraf=False):
    doc = Document()
    
    # SETUP MARGIN CUSTOM
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.20)
        section.right_margin = Cm(2.20)
        section.gutter = Cm(0)
    
    p_judul = doc.add_paragraph()
    p_judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_judul = p_judul.add_run("KERANGKA ACUAN KERJA/TERM OF REFERENCE\n")
    r_judul.bold = True
    r_judul.font.size = Pt(12)
    
    r_subjudul = p_judul.add_run(f"{meta['keg_title'].upper()}\n")
    r_subjudul.bold = True
    r_subjudul.font.size = Pt(12)
    
    table = doc.add_table(rows=0, cols=3)
    for row in table.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(0.2)
        row.cells[2].width = Inches(4.0)

    def add_meta_row(label, value):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = ":"
        cells[2].text = str(value)
        
    add_meta_row("Kementerian Negara/Lembaga", "(023) Kementerian Pendidikan, Kebudayaan, Riset dan Teknologi")
    add_meta_row("Unit Eselon I", "(17) Direktorat Jenderal Pendidikan Tinggi, Riset dan Teknologi")
    add_meta_row("Program", f"{meta['ro_code']} ({meta['ro_name']} ({meta['sumber_dana']}))")
    add_meta_row("Sasaran Program", meta['sasaran_prog'])
    add_meta_row("Indikator Kinerja Program", meta['ikp'])
    add_meta_row("Kegiatan", meta['kegiatan_induk'])
    add_meta_row("Sasaran Kegiatan", meta['sasaran_keg'])
    add_meta_row("Indikator Kinerja Kegiatan", meta['ikk'])
    add_meta_row("Klasifikasi Rincian Output", meta['kro_teks'])
    add_meta_row("Indikator KRO", meta['ind_kro'])
    add_meta_row("Rincian Output", meta['ro_teks'])
    add_meta_row("Volume RO", f"{meta['vol']}")
    add_meta_row("Satuan RO", meta['sat'])

    doc.add_paragraph("A. Latar Belakang").bold = True
    doc.add_paragraph("1. Dasar Hukum").bold = True
    
    dh_data = narasi.get('dasar_hukum', [])
    if isinstance(dh_data, list):
        for item in dh_data:
            p_bullet = doc.add_paragraph()
            p_bullet.paragraph_format.left_indent = Cm(0.7)
            p_bullet.add_run(f"•  {item}")
    else:
        for item in str(dh_data).split('\n'):
            p_bullet = doc.add_paragraph()
            p_bullet.paragraph_format.left_indent = Cm(0.7)
            p_bullet.add_run(f"•  {item.strip()}")
    
    doc.add_paragraph("2. Gambaran Umum").bold = True
    doc.add_paragraph(narasi.get('gambaran_umum', ''))
    
    doc.add_paragraph("B. Penerima Manfaat").bold = True
    doc.add_paragraph(narasi.get('penerima_manfaat', ''))
    
    doc.add_paragraph("C. Strategi Pencapaian Keluaran").bold = True
    doc.add_paragraph("1. Metode Pelaksanaan").bold = True
    doc.add_paragraph(narasi.get('metode_pelaksanaan', ''))
    
    doc.add_paragraph("2. Tahapan dan Waktu Pelaksanaan").bold = True
    doc.add_paragraph(narasi.get('tahapan_waktu', ''))
    
    doc.add_paragraph("D. Biaya Yang Diperlukan").bold = True
    doc.add_paragraph(narasi.get('biaya_diperlukan', ''))

    # TANDA TANGAN 
    doc.add_paragraph("\n")
    p_ttd = doc.add_paragraph()
    p_ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ttd.add_run(f"Samarinda, {meta['tgl_cetak']}\n")
    p_ttd.add_run(f"Dekan\n\n\n\n")
    
    r_nama = p_ttd.add_run(f"{meta['ketua']}\n")
    r_nama.bold = True
    if meta['nip']:
        r_nip = p_ttd.add_run(f"NIP {meta['nip']}")

    # TABEL PARAF DOCX (KIRI BAWAH)
    if tampilkan_paraf:
        p_paraf_title = doc.add_paragraph()
        table_paraf = doc.add_table(rows=4, cols=3)
        table_paraf.style = 'Table Grid'
        
        for row in table_paraf.rows:
            row.cells[0].width = Cm(1.2)
            row.cells[1].width = Cm(4.5)
            row.cells[2].width = Cm(3.0)
            
        hdr_cells = table_paraf.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[1].text = 'Jabatan'
        hdr_cells[2].text = 'Paraf'
        
        for i in range(1, 4):
            table_paraf.rows[i].cells[0].text = str(i)
            table_paraf.rows[i].cells[1].text = ""
            table_paraf.rows[i].cells[2].text = "\n\n" # Spasi agar baris lebih tinggi

    output = BytesIO()
    doc.save(output)
    return output.getvalue()

# --- BUILDER PDF / HTML PRINT-READY ---
def generate_tor_html(meta, narasi, tampilkan_paraf=False):
    dh_data = narasi.get('dasar_hukum', [])
    if isinstance(dh_data, list):
        dh_html = "<ul style='margin-top: 0; padding-left: 20px;'>" + "".join(f"<li>{item}</li>" for item in dh_data) + "</ul>"
    else:
        dh_html = str(dh_data).replace('\n', '<br>')

    nip_html = f"<br>NIP {meta['nip']}" if meta['nip'] else ""

    # STRING TABEL PARAF HTML
    paraf_html = ""
    if tampilkan_paraf:
        paraf_html = """
        <table style="width: 250px; border-collapse: collapse; float: left; margin-top: 60px; font-size: 8.5pt;">
            <tr>
                <th style="border: 1px solid black; padding: 4px; text-align: center; width: 15%;">No</th>
                <th style="border: 1px solid black; padding: 4px; text-align: center; width: 55%;">Jabatan</th>
                <th style="border: 1px solid black; padding: 4px; text-align: center; width: 30%;">Paraf</th>
            </tr>
            <tr><td style="border: 1px solid black; height: 30px; text-align: center; vertical-align: middle;">1</td><td style="border: 1px solid black;"></td><td style="border: 1px solid black;"></td></tr>
            <tr><td style="border: 1px solid black; height: 30px; text-align: center; vertical-align: middle;">2</td><td style="border: 1px solid black;"></td><td style="border: 1px solid black;"></td></tr>
            <tr><td style="border: 1px solid black; height: 30px; text-align: center; vertical-align: middle;">3</td><td style="border: 1px solid black;"></td><td style="border: 1px solid black;"></td></tr>
        </table>
        """

    html = f"""
    <!DOCTYPE html>
    <html><head><meta charset="utf-8">
    <style>
        @page {{ size: A4 portrait; margin: 20mm 22mm 20mm 22mm; }}
        body {{ font-family: 'Arial', sans-serif; font-size: 11pt; line-height: 1.5; color: #000; text-align: justify; }}
        .center {{ text-align: center; font-weight: bold; font-size: 12pt; margin-bottom: 5px; }}
        .subtitle {{ text-align: center; font-weight: bold; font-size: 12pt; margin-bottom: 20px; text-transform: uppercase; }}
        table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; }}
        td {{ padding: 2px 4px; vertical-align: top; }}
        .label {{ width: 35%; }}
        .titik {{ width: 3%; text-align: center; }}
        .value {{ width: 62%; }}
        .bab-title {{ font-weight: bold; margin-top: 15px; margin-bottom: 5px; }}
        .sub-bab {{ font-weight: bold; margin-top: 10px; margin-bottom: 5px; padding-left: 15px; }}
        .isi-text {{ margin-top: 0px; margin-bottom: 10px; padding-left: 15px; white-space: pre-wrap; }}
        .isi-sub-text {{ margin-top: 0px; margin-bottom: 10px; padding-left: 30px; white-space: pre-wrap; }}
        .ttd-box {{ width: 250px; float: right; text-align: left; margin-top: 40px; }}
    </style></head><body>
    
    <div class="center">KERANGKA ACUAN KERJA/TERM OF REFERENCE</div>
    <div class="subtitle">{meta['keg_title']}</div>
    
    <table>
        <tr><td class="label">Kementerian Negara/Lembaga</td><td class="titik">:</td><td class="value">(023) Kementerian Pendidikan, Kebudayaan, Riset dan Teknologi</td></tr>
        <tr><td class="label">Unit Eselon I</td><td class="titik">:</td><td class="value">(17) Direktorat Jenderal Pendidikan Tinggi, Riset dan Teknologi</td></tr>
        <tr><td class="label">Program</td><td class="titik">:</td><td class="value">{meta['ro_code']} ({meta['ro_name']} ({meta['sumber_dana']}))</td></tr>
        <tr><td class="label">Sasaran Program</td><td class="titik">:</td><td class="value">{meta['sasaran_prog']}</td></tr>
        <tr><td class="label">Indikator Kinerja Program</td><td class="titik">:</td><td class="value">{meta['ikp']}</td></tr>
        <tr><td class="label">Kegiatan</td><td class="titik">:</td><td class="value">{meta['kegiatan_induk']}</td></tr>
        <tr><td class="label">Sasaran Kegiatan</td><td class="titik">:</td><td class="value">{meta['sasaran_keg']}</td></tr>
        <tr><td class="label">Indikator Kinerja Kegiatan</td><td class="titik">:</td><td class="value">{meta['ikk']}</td></tr>
        <tr><td class="label">Klasifikasi Rincian Output</td><td class="titik">:</td><td class="value">{meta['kro_teks']}</td></tr>
        <tr><td class="label">Indikator KRO</td><td class="titik">:</td><td class="value">{meta['ind_kro']}</td></tr>
        <tr><td class="label">Rincian Output</td><td class="titik">:</td><td class="value">{meta['ro_teks']}</td></tr>
        <tr><td class="label">Volume RO</td><td class="titik">:</td><td class="value">{meta['vol']}</td></tr>
        <tr><td class="label">Satuan RO</td><td class="titik">:</td><td class="value">{meta['sat']}</td></tr>
    </table>
    
    <div class="bab-title">A. Latar Belakang</div>
    <div class="sub-bab">1. Dasar Hukum</div>
    <div class="isi-sub-text">{dh_html}</div>
    <div class="sub-bab">2. Gambaran Umum</div>
    <div class="isi-sub-text">{narasi.get('gambaran_umum', '')}</div>
    
    <div class="bab-title">B. Penerima Manfaat</div>
    <div class="isi-text">{narasi.get('penerima_manfaat', '')}</div>
    
    <div class="bab-title">C. Strategi Pencapaian Keluaran</div>
    <div class="sub-bab">1. Metode Pelaksanaan</div>
    <div class="isi-sub-text">{narasi.get('metode_pelaksanaan', '')}</div>
    <div class="sub-bab">2. Tahapan dan Waktu Pelaksanaan</div>
    <div class="isi-sub-text">{narasi.get('tahapan_waktu', '')}</div>
    
    <div class="bab-title">D. Biaya Yang Diperlukan</div>
    <div class="isi-text">{narasi.get('biaya_diperlukan', '')}</div>
    
    <div class="ttd-box">
        Samarinda, {meta['tgl_cetak']}<br>
        Dekan<br><br><br><br><br>
        <b>{meta['ketua']}</b>{nip_html}
    </div>
    
    {paraf_html}
    <div style="clear: both;"></div>
    
    </body></html>
    """
    return html

# =====================================================================
# TAMPILAN HALAMAN UTAMA
# =====================================================================
def show_page():
    st.title("🤖 Asisten Penyusun TOR Otomatis")
    st.caption("Didukung oleh Google Gemini AI. Menghasilkan struktur TOR standar Universitas.")
    
    df_utama, df_detail = load_active_rab()
    
    if df_utama.empty:
        st.warning("⚠️ Belum ada dokumen RAB yang berstatus AKTIF. Silakan aktifkan RAB di modul Pengolah RAB terlebih dahulu.")
        return

    if 'tor_json' not in st.session_state: st.session_state.tor_json = None
    
    kegiatan_list = sorted(df_utama['Kegiatan'].unique())
    
    tab_setup, tab_drafting, tab_cetak = st.tabs(["1️⃣ Setup Data Meta", "2️⃣ Drafting AI Gemini", "3️⃣ Finalisasi & Cetak"])

    # --- TAB 1: SETUP ---
    with tab_setup:
        st.subheader("Pilih Data Kegiatan (RAB Aktif)")
        pilih_keg = st.selectbox("Kegiatan:", kegiatan_list, format_func=lambda x: x.title())
        
        df_keg_utama = df_utama[df_utama['Kegiatan'] == pilih_keg].iloc[0]
        df_keg_det = df_detail[df_detail['ID_RAB'] == df_keg_utama['ID_RAB']]
        tot_rp = df_keg_det['Total_Biaya'].sum()
        
        st.info(f"**Sasaran:** {df_keg_utama['Sasaran']}\n\n**Total Pagu:** Rp {format_rupiah(tot_rp)}")
        
        st.markdown("---")
        st.subheader("Informasi Operasional & Penanggung Jawab")
        
        col1, col2, col3 = st.columns(3)
        in_ketua = col1.text_input("Nama Pejabat", value="Prof. Dr. M. Bahri Arifin, M.Hum.")
        in_nip = col2.text_input("NIP", value="196211271989031004")
        
        tgl_default = format_tgl_indo(df_keg_utama['Tanggal'])
        in_tgl_cetak = col3.text_input("Tanggal Cetak", value=tgl_default)
        
        in_poin = st.text_area("Catatan Latar Belakang untuk AI (Opsional)", placeholder="Ketik ide/alasan singkat mengapa kegiatan ini butuh dilaksanakan agar AI bisa merangkainya.")

    kro_code, kro_name = split_kode(df_keg_utama['KRO'])
    ro_code, ro_name = split_kode(df_keg_utama['RO'])
    
    meta_tor = {
        'sumber_dana': df_keg_utama['Sumber_Dana'],
        'ro_code': ro_code,
        'ro_name': ro_name,
        'sasaran_prog': f"Peningkatan {kro_name}",
        'ikp': f"Meningkatnya Kualitas {kro_name}",
        'kegiatan_induk': "7730. (Peningkatan Kualitas dan Kapasitas Perguruan Tinggi)",
        'sasaran_keg': df_keg_utama['Sasaran'],
        'ikk': "Meningkatnya Kualitas dan Kapasitas Perguruan Tinggi",
        'kro_teks': f"{kro_code}. ({kro_name})",
        'ind_kro': f"Meningkatnya Kualitas {ro_name}",
        'ro_teks': f"{ro_code}. ({ro_name})",
        'vol': df_keg_utama['Volume'],
        'sat': df_keg_utama['Satuan'],
        'ketua': in_ketua,
        'nip': in_nip, 
        'tgl_cetak': in_tgl_cetak,
        'keg_title': pilih_keg.title()
    }

    # --- TAB 2: AI DRAFTING ---
    with tab_drafting:
        st.subheader("Penyusunan Narasi Berbantu AI")
        st.markdown("AI Gemini akan membaca rincian RAB dan menyusun narasi TOR per bab.")
        
        if st.button("✨ Auto-Generate Narasi TOR", type="primary"):
            with st.spinner("AI sedang menganalisis data dan merangkai paragraf..."):
                list_barang = ", ".join(df_keg_det['Uraian'].tolist()[:15])
                
                hasil_json = generate_narasi_tor_json(
                    kegiatan=meta_tor['keg_title'],
                    total_anggaran=format_rupiah(tot_rp),
                    sasaran=meta_tor['sasaran_keg'],
                    list_belanja=list_barang,
                    poin_tambahan=in_poin
                )
                
                if hasil_json:
                    st.session_state.tor_json = hasil_json
                    st.success("Draft berhasil disusun! Silakan periksa dan edit di bawah ini.")
                
        if st.session_state.tor_json:
            st.markdown("### 📝 Editor Draft TOR")
            st.caption("Anda dapat menyunting langsung teks di bawah ini sebelum dicetak.")
            
            dh_val = st.session_state.tor_json.get('dasar_hukum', [])
            if isinstance(dh_val, list):
                dh_val = "\n".join(dh_val)
            
            with st.form("form_edit_tor"):
                edit_dh = st.text_area("A.1. Dasar Hukum (Tiap baris akan jadi 1 bullet)", value=dh_val, height=150)
                edit_gu = st.text_area("A.2. Gambaran Umum", value=st.session_state.tor_json.get('gambaran_umum', ''), height=200)
                edit_pm = st.text_area("B. Penerima Manfaat", value=st.session_state.tor_json.get('penerima_manfaat', ''), height=150)
                edit_mp = st.text_area("C.1. Metode Pelaksanaan", value=st.session_state.tor_json.get('metode_pelaksanaan', ''), height=150)
                edit_tw = st.text_area("C.2. Tahapan dan Waktu Pelaksanaan", value=st.session_state.tor_json.get('tahapan_waktu', ''), height=150)
                edit_bd = st.text_area("D. Biaya Yang Diperlukan", value=st.session_state.tor_json.get('biaya_diperlukan', ''), height=150)
                
                if st.form_submit_button("Simpan Perubahan Draft", type="primary"):
                    dh_list = [x.strip() for x in edit_dh.split('\n') if x.strip()]
                    
                    st.session_state.tor_json = {
                        'dasar_hukum': dh_list, 
                        'gambaran_umum': edit_gu, 
                        'penerima_manfaat': edit_pm,
                        'metode_pelaksanaan': edit_mp, 
                        'tahapan_waktu': edit_tw, 
                        'biaya_diperlukan': edit_bd
                    }
                    st.success("Perubahan tersimpan! Lanjut ke tab 3 untuk mencetak.")

    # --- TAB 3: CETAK ---
    with tab_cetak:
        st.subheader("Cetak Dokumen Resmi")
        if st.session_state.tor_json:
            st.success("Narasi telah siap. Metadata akan dilampirkan otomatis ke dalam dokumen.")
            
            # --- TOGGLE PARAF DITAMBAHKAN DI SINI ---
            tampilkan_paraf = st.checkbox("Tampilkan Tabel Paraf (Khusus Arsip Hardcopy Internal)", value=False)
            st.markdown("---")
            
            col_x1, col_x2 = st.columns(2)
            
            with col_x1:
                file_word = build_docx(meta_tor, st.session_state.tor_json, tampilkan_paraf)
                st.download_button(
                    label="📥 Download TOR (.docx)",
                    data=file_word,
                    file_name=f"TOR_{meta_tor['keg_title'].replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            with col_x2:
                html_print = generate_tor_html(meta_tor, st.session_state.tor_json, tampilkan_paraf)
                st.download_button(
                    label="📑 Cetak PDF (HTML Print-Ready)",
                    data=html_print.encode('utf-8'),
                    file_name=f"TOR_{meta_tor['keg_title'].replace(' ', '_')}.html",
                    mime="text/html",
                    use_container_width=True,
                    help="Buka file di browser Chrome/Edge, lalu tekan Ctrl+P. Pilih orientasi Portrait."
                )
        else:
            st.warning("⚠️ Anda belum menyusun narasi. Silakan klik 'Auto-Generate Narasi TOR' di tab 2.")
