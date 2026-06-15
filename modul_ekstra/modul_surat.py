import streamlit as st
import google.generativeai as genai
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import json
import os
import base64
from datetime import datetime

from utils import log_audit, format_tgl_indo, engine

# =======================================================
# FUNGSI DATABASE UNTUK DRAF PERMANEN
# =======================================================
def load_draft_from_db(username):
    try:
        with engine.connect() as conn:
            df = pd.read_sql("SELECT * FROM rab_draft_surat WHERE \"Username\"='" + username + "'", conn)
            if not df.empty:
                return json.loads(df.iloc[0]['Meta_JSON']), json.loads(df.iloc[0]['Narasi_JSON'])
    except Exception as e:
        err_str = str(e).lower()
        if "does not exist" in err_str or "not found" in err_str or "relation" in err_str:
            df_kosong = pd.DataFrame(columns=["Username", "Meta_JSON", "Narasi_JSON", "Waktu_Simpan"])
            try:
                with engine.begin() as conn:
                    df_kosong.to_sql("rab_draft_surat", conn, if_exists="append", index=False)
            except: pass
    return None, None

def save_draft_to_db(username, meta, narasi):
    try:
        with engine.connect() as conn:
            df_all = pd.read_sql("SELECT * FROM rab_draft_surat", conn)
    except:
        df_all = pd.DataFrame(columns=["Username", "Meta_JSON", "Narasi_JSON", "Waktu_Simpan"])
        
    df_all = df_all[df_all["Username"] != username]
    new_row = pd.DataFrame([{
        "Username": username,
        "Meta_JSON": json.dumps(meta),
        "Narasi_JSON": json.dumps(narasi),
        "Waktu_Simpan": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }])
    df_all = pd.concat([df_all, new_row], ignore_index=True)
    
    try:
        with engine.begin() as conn:
            df_all.to_sql("rab_draft_surat", conn, if_exists="replace", index=False)
    except Exception as e:
        st.error("Gagal menyimpan draf ke database: " + str(e))

def delete_draft_from_db(username):
    try:
        with engine.connect() as conn:
            df_all = pd.read_sql("SELECT * FROM rab_draft_surat", conn)
        df_all = df_all[df_all["Username"] != username]
        with engine.begin() as conn:
            df_all.to_sql("rab_draft_surat", conn, if_exists="replace", index=False)
    except: pass

# =======================================================
# FUNGSI BANTUAN & AI DENGAN LOGIKA TEMPLATE
# =======================================================
def get_image_base64(image_path):
    if os.path.exists(image_path):
        with open(image_path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode('utf-8')
    return ""

def generate_surat_ai(jenis, hal, tujuan, isi_poin):
    prompt_template = """Anda adalah staf administrasi Fakultas Ilmu Budaya Universitas Mulawarman. 
Buatlah draf narasi untuk [JENIS_SURAT] berdasarkan data:
- Hal / Perihal: [HAL]
- Tujuan: [TUJUAN]
- Poin Utama Isi Surat: [POIN]

ATURAN STRUKTUR KHUSUS UNTUK [JENIS_SURAT]:
[ATURAN_KHUSUS]

Gunakan bahasa Indonesia baku dan formal khas instansi pemerintah.
Output JSON murni dengan kunci: "pembuka", "isi", "penutup". Dilarang menggunakan markdown block seperti ```json atau 
```. Hasilkan teks mentah yang bisa langsung diparse.
"""
    
    aturan = ""
    if jenis == "Surat Pengantar / Biasa":
        aturan = "1. Pembuka: Pengantar tujuan pengiriman surat.\n2. Isi: Penjelasan maksud dan tujuan sesuai poin utama.\n3. Penutup: Demikian disampaikan, terima kasih."
    elif jenis == "Surat Undangan Rapat / Acara":
        aturan = "1. Pembuka: Salam hormat dan tujuan undangan.\n2. Isi: Rincian acara (hari, tanggal, waktu, tempat, agenda) yang diekstrak dari poin utama. Buat dalam format naratif yang rapi.\n3. Penutup: Harapan kehadiran tepat waktu dan ucapan terima kasih."
    elif jenis == "Surat Pemberitahuan / Edaran":
        aturan = "1. Pembuka: Dasar atau latar belakang pemberitahuan.\n2. Isi: Informasi utama yang perlu diketahui khalayak/tujuan. Jabarkan dengan terstruktur.\n3. Penutup: Himbauan untuk mematuhi/mengetahui dan ucapan terima kasih."
    elif jenis == "Surat Permohonan (Dana / Izin)":
        aturan = "1. Pembuka: Latar belakang permohonan.\n2. Isi: Rincian hal yang dimohonkan dengan bahasa birokrasi yang persuasif namun formal.\n3. Penutup: Harapan agar permohonan dikabulkan dan ucapan terima kasih."
    elif jenis == "Surat Tugas":
        aturan = "1. Pembuka: Dasar penugasan (jika ada di poin utama) atau langsung ke pernyataan menugaskan.\n2. Isi: Sebutkan untuk melaksanakan tugas apa, beserta detail waktu/tempat.\n3. Penutup: Perintah agar tugas dilaksanakan dengan penuh tanggung jawab."
    
    prompt = prompt_template.replace("[JENIS_SURAT]", jenis)
    prompt = prompt.replace("[HAL]", hal)
    prompt = prompt.replace("[TUJUAN]", tujuan)
    prompt = prompt.replace("[POIN]", isi_poin)
    prompt = prompt.replace("[ATURAN_KHUSUS]", aturan)
    
    try:
        genai.configure(api_key=st.secrets["GEMINI_API_KEY_NEW"])
        model_list = genai.list_models()
        model_yang_bisa = [m.name for m in model_list if 'generateContent' in m.supported_generation_methods]
        
        if not model_yang_bisa: return None
        model_pilihan = model_yang_bisa[0]
        for m in model_yang_bisa:
            if 'gemini-1.5' in m:
                model_pilihan = m
                break
        
        model = genai.GenerativeModel(model_pilihan.replace("models/", ""))
        respons = model.generate_content(prompt)
        
        clean_text = respons.text.strip()
        if clean_text.startswith("```json"): clean_text = clean_text[7:]
        elif clean_text.startswith("```"): clean_text = clean_text[3:]
        if clean_text.endswith("```"): clean_text = clean_text[:-3]
        
        return json.loads(clean_text.strip())
    
    except Exception as e:
        if "429" in str(e):
            st.error("❌ Kuota AI Habis (Limit). Tunggu 20-30 detik lalu klik Generate lagi.")
        else:
            st.error("Error AI saat memproses JSON. Harap klik Generate ulang. Detail: " + str(e))
        return None

# =======================================================
# FUNGSI CETAK DOCX & HTML (DENGAN OPSIONAL LAMPIRAN)
# =======================================================
def build_surat_docx(meta, narasi, tampilkan_paraf=True, df_lampiran=None):
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

    if os.path.exists("Header Kop Surat.jpg"):
        p_kop = doc.add_paragraph()
        p_kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_kop.add_run().add_picture("Header Kop Surat.jpg", width=Cm(16))
    else:
        p_kop = doc.add_paragraph()
        p_kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p_kop.add_run("KEMENTERIAN PENDIDIKAN TINGGI, SAINS, DAN TEKNOLOGI\nUNIVERSITAS MULAWARMAN\nFAKULTAS ILMU BUDAYA\n")
        r1.bold = True
        p_kop.add_run("Jl. Ki Hajar Dewantara, Kampus Gunung Kelua, Samarinda 75123\nTelepon (0541) 7809033")
    
    doc.add_paragraph()
    
    t_meta = doc.add_table(rows=3, cols=4)
    for row in t_meta.rows:
        row.cells[0].width = Cm(2.0); row.cells[1].width = Cm(0.5); row.cells[2].width = Cm(8.0); row.cells[3].width = Cm(5.0)

    t_meta.cell(0, 0).text = "Nomor"; t_meta.cell(0, 1).text = ":"
    t_meta.cell(0, 2).text = meta['nomor']; t_meta.cell(0, 3).text = meta['tgl_surat']
    t_meta.cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    t_meta.cell(1, 0).text = "Lampiran"; t_meta.cell(1, 1).text = ":"; t_meta.cell(1, 2).text = meta['lampiran']
    t_meta.cell(2, 0).text = "Hal"; t_meta.cell(2, 1).text = ":"; t_meta.cell(2, 2).text = meta['hal']

    p_tujuan = doc.add_paragraph()
    p_tujuan.add_run("\nYth. " + meta['tujuan'] + "\nSamarinda\n")

    p_buka = doc.add_paragraph(narasi.get('pembuka', ''))
    p_buka.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_isi = doc.add_paragraph(narasi.get('isi', ''))
    p_isi.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_tutup = doc.add_paragraph(narasi.get('penutup', ''))
    p_tutup.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_ttd = doc.add_paragraph()
    p_ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ttd.add_run("Dekan,\n\n\n\n")
    r_ttd = p_ttd.add_run(meta['dekan_nama'] + "\nNIP " + meta['dekan_nip'])
    r_ttd.bold = True
    
    doc.add_paragraph("\n")
    
    if tampilkan_paraf:
        t_paraf = doc.add_table(rows=4, cols=3)
        t_paraf.style = 'Table Grid'
        for row in t_paraf.rows:
            row.cells[0].width = Cm(1.0); row.cells[1].width = Cm(6.5); row.cells[2].width = Cm(2.5)

        data = [["NO", "JABATAN", "PARAF"], ["1", "Wakil Dekan Bidang Keuangan dan Umum", ""], ["2", "Kepala Bagian Umum", ""], ["3", "Staf Perencanaan", ""]]
        for i, row in enumerate(t_paraf.rows):
            for j, cell in enumerate(row.cells): 
                cell.text = data[i][j]
                if i == 0: cell.paragraphs[0].runs[0].bold = True

    # --- INJEKSI HALAMAN LAMPIRAN (OPSIONAL) ---
    if df_lampiran is not None and not df_lampiran.empty:
        doc.add_page_break()
        p_lamp = doc.add_paragraph()
        p_lamp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_lamp = p_lamp.add_run("LAMPIRAN SURAT\n")
        r_lamp.bold = True
        p_lamp.add_run("Nomor\t: " + meta['nomor'] + "\nHal\t\t: " + meta['hal'] + "\n")
        
        t_lamp = doc.add_table(rows=1, cols=len(df_lampiran.columns))
        t_lamp.style = 'Table Grid'
        
        for j, col_name in enumerate(df_lampiran.columns):
            t_lamp.cell(0, j).text = str(col_name)
            t_lamp.cell(0, j).paragraphs[0].runs[0].bold = True
            
        for _, row_data in df_lampiran.iterrows():
            row_cells = t_lamp.add_row().cells
            for j, val in enumerate(row_data):
                row_cells[j].text = str(val)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def generate_surat_html(meta, narasi, tampilkan_paraf=True, df_lampiran=None):
    img_header = get_image_base64("Header Kop Surat.jpg")
    img_footer = get_image_base64("Footer Kop Surat.jpg")
    img_maskot = get_image_base64("Maskot Baru.png")
    
    html_header = '<img src="data:image/jpeg;base64,' + img_header + '" class="header-img">' if img_header else '<h2 style="text-align:center;">[GAMBAR HEADER TIDAK DITEMUKAN]</h2><hr>'
    html_footer = '<img src="data:image/jpeg;base64,' + img_footer + '" class="footer-img">' if img_footer else ''
    html_watermark = '<img src="data:image/png;base64,' + img_maskot + '" class="watermark-img">' if img_maskot else ''

    html_paraf = ""
    if tampilkan_paraf:
        html_paraf = """
        <div class="paraf-box">
            <table class="paraf-table">
                <tr><th style="width: 10%; text-align: center;">NO</th><th style="width: 65%;">JABATAN</th><th style="width: 25%;">PARAF</th></tr>
                <tr><td style="text-align: center;">1</td><td>Wakil Dekan Bidang Keuangan dan Umum</td><td></td></tr>
                <tr><td style="text-align: center;">2</td><td>Kepala Bagian Umum</td><td></td></tr>
                <tr><td style="text-align: center;">3</td><td>Staf Perencanaan</td><td></td></tr>
            </table>
        </div>
        """

    # --- INJEKSI HTML LAMPIRAN (OPSIONAL) ---
    html_lampiran_tabel = ""
    if df_lampiran is not None and not df_lampiran.empty:
        tabel_data_html = df_lampiran.to_html(index=False, classes="lampiran-table", border=0)
        html_lampiran_tabel = '<div style="page-break-before: always;"></div>'
        html_lampiran_tabel += '<h4 style="margin-bottom: 5px;">LAMPIRAN SURAT</h4>'
        html_lampiran_tabel += '<table style="width: 100%; border-collapse: collapse; border: none; margin-bottom: 15px;">'
        html_lampiran_tabel += '<tr><td style="width: 12%; border: none; padding: 2px;">Nomor</td><td style="width: 2%; border: none; padding: 2px;">:</td><td style="border: none; padding: 2px;">' + meta['nomor'] + '</td></tr>'
        html_lampiran_tabel += '<tr><td style="border: none; padding: 2px;">Hal</td><td style="border: none; padding: 2px;">:</td><td style="border: none; padding: 2px;">' + meta['hal'] + '</td></tr>'
        html_lampiran_tabel += '</table>'
        html_lampiran_tabel += tabel_data_html

    html_template = """<!DOCTYPE html>
    <html><head><meta charset="utf-8">
    <style>
        @page { size: A4 portrait; margin: 10mm; }
        body { font-family: 'Times New Roman', Times, serif; font-size: 11.5pt; line-height: 1.5; color: #000; text-align: justify; position: relative; padding-bottom: 120px; }
        .header-img { width: 100%; display: block; margin-bottom: 25px; }
        .watermark-img { position: fixed; top: 50%; left: 50%; transform: translate(-50%, -50%); opacity: 0.15; width: 450px; z-index: -1; }
        .footer-img { position: fixed; bottom: 0; left: 0; width: 100%; z-index: 10; }
        table.meta-table { width: 100%; border-collapse: collapse; margin-bottom: 15px; border: none; }
        table.meta-table td { padding: 1px; vertical-align: top; border: none; }
        .isi-surat p { text-indent: 0; margin-top: 5px; margin-bottom: 15px; text-align: justify; white-space: pre-line; }
        .tujuan { margin-bottom: 20px; margin-top: 15px; }
        .ttd-box { width: 250px; float: right; text-align: left; margin-top: 30px; }
        .paraf-box { width: 320px; float: left; margin-top: 80px; clear: left; }
        table.paraf-table { width: 100%; border-collapse: collapse; font-size: 9.5pt; font-family: 'Arial', sans-serif; }
        table.paraf-table th, table.paraf-table td { border: 1px solid black; padding: 4px 6px; }
        table.paraf-table th { text-align: left; }
        /* Style Khusus Tabel Lampiran */
        table.lampiran-table { width: 100%; border-collapse: collapse; font-size: 10pt; margin-top: 10px; }
        table.lampiran-table th, table.lampiran-table td { border: 1px solid black; padding: 5px; text-align: left; }
        table.lampiran-table th { background-color: #f2f2f2; }
    </style></head><body>
    
    [WATERMARK]
    [HEADER]
    
    <table class="meta-table">
        <tr><td style="width: 12%;">Nomor</td><td style="width: 2%;">:</td><td style="width: 56%;">[NOMOR]</td><td style="width: 30%; text-align: right;">[TGL_SURAT]</td></tr>
        <tr><td>Lampiran</td><td>:</td><td colspan="2">[LAMPIRAN]</td></tr>
        <tr><td>Hal</td><td>:</td><td colspan="2">[HAL]</td></tr>
    </table>
    
    <div class="tujuan">Yth. [TUJUAN]<br>Samarinda</div>
    
    <div class="isi-surat">
        <p>[PEMBUKA]</p>
        <p>[ISI]</p>
        <p>[PENUTUP]</p>
    </div>
    
    <div class="ttd-box">Dekan,<br><br><br><br><br><b>[DEKAN_NAMA]</b><br>NIP [DEKAN_NIP]</div>
    [PARAF]
    [FOOTER]
    
    <div style="clear: both;"></div>
    [LAMPIRAN_TABEL]
    </body></html>"""

    html_result = html_template.replace("[WATERMARK]", html_watermark).replace("[HEADER]", html_header)
    html_result = html_result.replace("[NOMOR]", meta['nomor']).replace("[TGL_SURAT]", meta['tgl_surat'])
    html_result = html_result.replace("[LAMPIRAN]", meta['lampiran']).replace("[HAL]", meta['hal'])
    html_result = html_result.replace("[TUJUAN]", meta['tujuan'].replace('\n', '<br>'))
    html_result = html_result.replace("[PEMBUKA]", narasi.get('pembuka', '')).replace("[ISI]", narasi.get('isi', '')).replace("[PENUTUP]", narasi.get('penutup', ''))
    html_result = html_result.replace("[DEKAN_NAMA]", meta['dekan_nama']).replace("[DEKAN_NIP]", meta['dekan_nip'])
    html_result = html_result.replace("[PARAF]", html_paraf).replace("[FOOTER]", html_footer)
    html_result = html_result.replace("[LAMPIRAN_TABEL]", html_lampiran_tabel)
    
    return html_result

# =======================================================
# ANTARMUKA PENGGUNA (UI)
# =======================================================
def show_page():
    username_aktif = st.session_state.get('username', 'user_anonim')
    
    if 'surat_draft_loaded' not in st.session_state:
        db_meta, db_narasi = load_draft_from_db(username_aktif)
        if db_meta and db_narasi:
            st.session_state['surat_meta'] = db_meta
            st.session_state['surat_narasi'] = db_narasi
            st.toast("💡 Memuat draf pekerjaan Anda sebelumnya...", icon="🔄")
        st.session_state['surat_draft_loaded'] = True
        
    def_meta = st.session_state.get('surat_meta', {})

    col_title1, col_title2 = st.columns([3, 1])
    col_title1.title("✉️ Pengolah Surat Otomatis")
    col_title1.caption("Sistem otomatis menyimpan draf Anda setiap kali disimulasikan atau disimpan.")
    
    if col_title2.button("🗑️ Reset & Buat Baru", use_container_width=True):
        delete_draft_from_db(username_aktif)
        if 'surat_meta' in st.session_state: del st.session_state['surat_meta']
        if 'surat_narasi' in st.session_state: del st.session_state['surat_narasi']
        st.success("Draf dibersihkan. Memulai halaman baru!"); st.rerun()

    missing_files = []
    if not os.path.exists("Header Kop Surat.jpg"): missing_files.append("Header Kop Surat.jpg")
    if not os.path.exists("Footer Kop Surat.jpg"): missing_files.append("Footer Kop Surat.jpg")
    if not os.path.exists("Maskot Baru.png"): missing_files.append("Maskot Baru.png")
    
    if missing_files:
        st.warning("⚠️ Peringatan: File template gambar belum lengkap: **" + ", ".join(missing_files) + "**. Cetakan HTML tidak akan memiliki gambar tersebut.")

    daftar_template = ["Surat Pengantar / Biasa", "Surat Undangan Rapat / Acara", "Surat Pemberitahuan / Edaran", "Surat Permohonan (Dana / Izin)", "Surat Tugas"]
    
    with st.container(border=True):
        st.subheader("1. Identitas Surat")
        
        idx_template = daftar_template.index(def_meta.get('jenis_surat', daftar_template[0])) if def_meta.get('jenis_surat') in daftar_template else 0
        jenis_surat_terpilih = st.selectbox("📋 Pilih Template Kerangka Surat:", daftar_template, index=idx_template)
        
        col1, col2, col3 = st.columns(3)
        meta = {
            'jenis_surat': jenis_surat_terpilih,
            'nomor': col1.text_input("Nomor Surat", def_meta.get('nomor', "499/UN17.13/PR.03.01/" + str(datetime.now().year))),
            'lampiran': col2.text_input("Lampiran", def_meta.get('lampiran', "1 (satu) berkas")),
            'tgl_surat': col3.text_input("Tanggal Surat", def_meta.get('tgl_surat', format_tgl_indo(datetime.now().strftime("%Y-%m-%d")))),
            'hal': st.text_input("Hal / Perihal", def_meta.get('hal', "Tindak Lanjut Permintaan Data")),
            'tujuan': st.text_area("Yth. / Pihak Tujuan", def_meta.get('tujuan', "Rektor\nUniversitas Mulawarman\nc.q Wakil Rektor Bidang Perencanaan, Kerjasama, dan Sistem Informasi")),
            'dekan_nama': "Prof. Dr. M. Bahri Arifin, M.Hum.",
            'dekan_nip': "196211271989031004"
        }
        
        st.markdown("---")
        opsi_paraf = st.checkbox("Tampilkan Kolom Paraf Pejabat", value=True, help="Hapus centang jika surat ini tidak membutuhkan paraf berjenjang.")
    
    with st.container(border=True):
        st.subheader("2. Materi & Poin Utama Surat")
        st.info("💡 AI akan mengemas teks ini menjadi struktur formal bergaya **" + jenis_surat_terpilih + "**.")
        isi_poin = st.text_area("Ketik intisari yang ingin Anda sampaikan:", def_meta.get('poin_mentah', "FIB mengirimkan Data Alokasi Anggaran Peningkatan Kualitas Penelitian dan Pengabdian kepada Masyarakat Tahun Anggaran 2025 dan 2026. Data ini sebagai bahan penyusunan rencana kerja."))
        meta['poin_mentah'] = isi_poin
        
        if st.button("✨ Generate Narasi Surat", type="primary"):
            with st.spinner("AI sedang merangkai bahasa dengan template " + jenis_surat_terpilih + "..."):
                narasi = generate_surat_ai(meta['jenis_surat'], meta['hal'], meta['tujuan'], isi_poin)
                if narasi:
                    st.session_state['surat_narasi'] = narasi
                    st.session_state['surat_meta'] = meta
                    save_draft_to_db(username_aktif, meta, narasi)
                    st.success("Draft narasi selesai disusun & otomatis di-Backup ke Database!")

    if 'surat_narasi' in st.session_state:
        st.markdown("---")
        st.subheader("3. Editor Draft & Cetak Dokumen")
        
        with st.form("form_edit_surat"):
            edit_pembuka = st.text_area("Paragraf Pembuka", value=st.session_state['surat_narasi'].get('pembuka', ''), height=100)
            edit_isi = st.text_area("Paragraf Isi", value=st.session_state['surat_narasi'].get('isi', ''), height=150)
            edit_penutup = st.text_area("Paragraf Penutup", value=st.session_state['surat_narasi'].get('penutup', ''), height=80)
            
            if st.form_submit_button("💾 Simpan Perubahan Teks"):
                st.session_state['surat_narasi'] = {'pembuka': edit_pembuka, 'isi': edit_isi, 'penutup': edit_penutup}
                st.session_state['surat_meta'] = meta
                save_draft_to_db(username_aktif, meta, st.session_state['surat_narasi'])
                st.success("Perubahan Teks dan Identitas Surat berhasil disimpan secara Permanen!")
        
        # --- FITUR UPLOAD LAMPIRAN ---
        st.markdown("#### 📎 Sisipkan Lampiran Tabel (Opsional)")
        file_lampiran = st.file_uploader("Upload File Excel/CSV jika surat ini memiliki lampiran tabel data:", type=["xlsx", "csv"])
        df_lampiran = None
        
        if file_lampiran:
            try:
                if file_lampiran.name.endswith(".csv"):
                    df_lampiran = pd.read_csv(file_lampiran)
                else:
                    df_lampiran = pd.read_excel(file_lampiran)
                st.success("Tabel lampiran berhasil dimuat! (" + str(len(df_lampiran)) + " baris data terdeteksi)")
                with st.expander("Lihat Pratinjau Tabel Lampiran"):
                    st.dataframe(df_lampiran, use_container_width=True)
            except Exception as e:
                st.error("Gagal membaca file Excel/CSV: " + str(e))
        
        st.markdown("<br>", unsafe_allow_html=True)
        col_x1, col_x2 = st.columns(2)
        
        hal_safe = meta['hal'].replace(' ', '_').replace('/', '_')
        
        with col_x1:
            file_word = build_surat_docx(meta, st.session_state['surat_narasi'], tampilkan_paraf=opsi_paraf, df_lampiran=df_lampiran)
            st.download_button(
                label="📥 Download Surat (.docx)",
                data=file_word,
                file_name="Surat_" + hal_safe + ".docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        with col_x2:
            html_print = generate_surat_html(meta, st.session_state['surat_narasi'], tampilkan_paraf=opsi_paraf, df_lampiran=df_lampiran)
            if st.download_button(
                label="📑 Cetak PDF (Format Template Gambar)",
                data=html_print.encode('utf-8'),
                file_name="Surat_" + hal_safe + ".html",
                mime="text/html",
                use_container_width=True,
                help="Buka file ini di browser Chrome/Edge, lalu tekan Ctrl+P. Gambar Header, Footer, dan Maskot akan otomatis termuat."
            ):
                log_audit("BUAT SURAT", "Hal: " + meta['hal'] + " (Cetak Template: " + jenis_surat_terpilih + ")")

show_page()
